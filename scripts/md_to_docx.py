#!/usr/bin/env python3
"""
Markdown to DOCX Converter — General-purpose script.

Converts any Markdown file to a Microsoft Word (.docx) document,
preserving headings, tables, lists, bold/italic, code blocks,
inline code, images, and LaTeX math (rendered as native OMML equations).

Requires: python-docx, latex2mathml, lxml
Optional: mml2omml.xsl (bundled in scripts/ directory)

Usage:
    python scripts/md_to_docx.py <input.md> [output.docx] [--font SIZE] [--style STYLE]

Examples:
    python scripts/md_to_docx.py docs/Summary/docs_ch2-ch6.md
    python scripts/md_to_docx.py README.md output.docx --font 12
    python scripts/md_to_docx.py report.md report.docx --style modern
"""

import argparse
import re
import subprocess
import sys
import tempfile
from pathlib import Path

try:
    from PIL import Image as PILImage
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)

try:
    from latex2mathml.converter import convert as latex_to_mathml
    from lxml import etree as lxml_etree
    HAS_LATEX_MATH = True
except ImportError:
    HAS_LATEX_MATH = False

try:
    import markdown as md_lib
except ImportError:
    md_lib = None


# ─── MathML → OMML converter ────────────────────────────────────────────────────

class MathConverter:
    """Convert LaTeX math to Office Math Markup Language (OMML) for DOCX."""

    def __init__(self, xsl_path: str = None):
        self._transform = None
        if not HAS_LATEX_MATH:
            return
        if xsl_path is None:
            xsl_path = Path(__file__).parent / "mml2omml.xsl"
        xsl_file = Path(xsl_path)
        if xsl_file.exists():
            xsl_doc = lxml_etree.parse(str(xsl_file))
            self._transform = lxml_etree.XSLT(xsl_doc)

    @property
    def available(self) -> bool:
        return self._transform is not None

    def latex_to_omml(self, latex: str) -> lxml_etree._Element | None:
        """Convert LaTeX string to OMML element tree. Returns None on failure."""
        if not self._transform:
            return None
        try:
            # Clean up LaTeX for latex2mathml compatibility
            latex_clean = latex.strip()
            # Remove \text{} wrapping that sometimes causes issues
            # but keep the content
            mathml_str = latex_to_mathml(latex_clean)
            mathml_tree = lxml_etree.fromstring(mathml_str.encode())
            omml_tree = self._transform(mathml_tree)
            return omml_tree.getroot()
        except Exception:
            return None

    def latex_to_omml_str(self, latex: str) -> str | None:
        """Convert LaTeX string to OMML XML string. Returns None on failure."""
        elem = self.latex_to_omml(latex)
        if elem is not None:
            return lxml_etree.tostring(elem, encoding="unicode")
        return None


# ─── Style presets ───────────────────────────────────────────────────────────

STYLE_CLASSIC = {
    "body_font": "Times New Roman",
    "heading_font": "Arial",
    "code_font": "Courier New",
    "body_size": 12,
    "line_spacing": 1.15,
    "heading_sizes": {1: 22, 2: 18, 3: 15, 4: 13, 5: 12, 6: 11},
    "heading_color": RGBColor(0x1A, 0x3C, 0x6E),
    "table_header_bg": RGBColor(0x1A, 0x3C, 0x6E),
    "table_header_fg": RGBColor(0xFF, 0xFF, 0xFF),
}

STYLE_MODERN = {
    "body_font": "Calibri",
    "heading_font": "Calibri",
    "code_font": "Consolas",
    "body_size": 11,
    "line_spacing": 1.3,
    "heading_sizes": {1: 24, 2: 20, 3: 16, 4: 14, 5: 12, 6: 11},
    "heading_color": RGBColor(0x00, 0x70, 0xC0),
    "table_header_bg": RGBColor(0x00, 0x70, 0xC0),
    "table_header_fg": RGBColor(0xFF, 0xFF, 0xFF),
}

STYLE_ACADEMIC = {
    "body_font": "Times New Roman",
    "heading_font": "Times New Roman",
    "code_font": "Courier New",
    "body_size": 12,
    "line_spacing": 1.5,
    "heading_sizes": {1: 20, 2: 16, 3: 14, 4: 13, 5: 12, 6: 11},
    "heading_color": RGBColor(0x00, 0x00, 0x00),
    "table_header_bg": RGBColor(0x40, 0x40, 0x40),
    "table_header_fg": RGBColor(0xFF, 0xFF, 0xFF),
}

STYLES = {"classic": STYLE_CLASSIC, "modern": STYLE_MODERN, "academic": STYLE_ACADEMIC}


# ─── Mermaid diagram renderer ─────────────────────────────────────────────────

class MermaidRenderer:
    """Render Mermaid diagrams to PNG images using mmdc CLI."""

    def __init__(self):
        self._mmdc = self._find_mmdc()
        self._tmp_dir = tempfile.mkdtemp(prefix="mermaid_")

    @staticmethod
    def _find_mmdc() -> str | None:
        """Find mmdc (mermaid-cli) executable."""
        import shutil
        return shutil.which("mmdc")

    @property
    def available(self) -> bool:
        return self._mmdc is not None

    def render(self, mermaid_code: str, width: int = 1200) -> str | None:
        """Render mermaid code to a PNG file. Returns path to PNG or None on failure."""
        if not self._mmdc:
            return None

        # Write mermaid code to a temp .mmd file
        mmd_path = Path(self._tmp_dir) / f"diagram_{id(mermaid_code) % 100000}.mmd"
        png_path = mmd_path.with_suffix(".png")

        try:
            mmd_path.write_text(mermaid_code, encoding="utf-8")

            cmd = [
                self._mmdc,
                "-i", str(mmd_path),
                "-o", str(png_path),
                "-w", str(width),
                "-b", "white",
                "--quiet",
            ]

            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60,
            )

            if result.returncode != 0:
                print(f"  ⚠️ Mermaid render error: {result.stderr.strip()}", file=sys.stderr)
                return None

            if png_path.exists():
                return str(png_path)

            return None
        except subprocess.TimeoutExpired:
            print("  ⚠️ Mermaid render timed out", file=sys.stderr)
            return None
        except Exception as e:
            print(f"  ⚠️ Mermaid render failed: {e}", file=sys.stderr)
            return None


# ─── Markdown parser ─────────────────────────────────────────────────────────

class MarkdownParser:
    """Parse markdown into structured blocks for DOCX rendering."""

    def __init__(self, text: str):
        self.text = text
        self.blocks: list[dict] = []
        self._parse()

    def _parse(self):
        lines = self.text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # ── Code block ────────────────────────────────────────────────
            if line.strip().startswith("```"):
                lang = line.strip()[3:].strip()
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                content = "\n".join(code_lines)
                # Detect mermaid diagrams
                if lang.lower() == "mermaid":
                    self.blocks.append({"type": "mermaid", "content": content})
                else:
                    self.blocks.append({"type": "code", "lang": lang, "content": content})
                i += 1
                continue

            # ── LaTeX math block ($$...$$) ────────────────────────────────
            if line.strip().startswith("$$"):
                math_lines = [line.strip()[2:]]  # strip opening $$
                if line.strip().endswith("$$") and len(line.strip()) > 4:
                    # Single-line $$...$$
                    content = line.strip()[2:-2].strip()
                    self.blocks.append({"type": "math_block", "content": content})
                    i += 1
                    continue
                i += 1
                while i < len(lines) and not lines[i].strip().endswith("$$"):
                    math_lines.append(lines[i])
                    i += 1
                if i < len(lines):
                    last = lines[i].strip()
                    if last.endswith("$$"):
                        math_lines.append(last[:-2])
                content = "\n".join(math_lines).strip()
                self.blocks.append({"type": "math_block", "content": content})
                i += 1
                continue

            # ── Table ─────────────────────────────────────────────────────
            if "|" in line and line.strip().startswith("|"):
                table_lines = []
                while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                self._parse_table(table_lines)
                continue

            # ── Heading ───────────────────────────────────────────────────
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                self.blocks.append({"type": "heading", "level": level, "text": text})
                i += 1
                continue

            # ── Horizontal rule ────────────────────────────────────────────
            if re.match(r"^(\*{3,}|-{3,}|_{3,})\s*$", line.strip()):
                self.blocks.append({"type": "hr"})
                i += 1
                continue

            # ── Unordered list ─────────────────────────────────────────────
            if re.match(r"^\s*[-*+]\s+", line):
                list_items = []
                while i < len(lines):
                    m = re.match(r"^(\s*)[-*+]\s+(.+)$", lines[i])
                    if not m:
                        break
                    indent = len(m.group(1)) // 2
                    list_items.append({"text": m.group(2), "indent": indent})
                    i += 1
                self.blocks.append({"type": "ul", "items": list_items})
                continue

            # ── Ordered list ──────────────────────────────────────────────
            if re.match(r"^\s*\d+\.\s+", line):
                list_items = []
                while i < len(lines):
                    m = re.match(r"^(\s*)\d+\.\s+(.+)$", lines[i])
                    if not m:
                        break
                    indent = len(m.group(1)) // 2
                    list_items.append({"text": m.group(2), "indent": indent})
                    i += 1
                self.blocks.append({"type": "ol", "items": list_items})
                continue

            # ── Empty line ────────────────────────────────────────────────
            if not line.strip():
                i += 1
                continue

            # ── Paragraph (collect consecutive lines) ─────────────────────
            para_lines = []
            while i < len(lines):
                ln = lines[i]
                # Stop if we hit a structural element
                if not ln.strip():
                    break
                if ln.strip().startswith("```") or ln.strip().startswith("$$"):
                    break
                if re.match(r"^#{1,6}\s+", ln):
                    break
                if re.match(r"^[-*+]\s+", ln.strip()):
                    break
                if re.match(r"^\d+\.\s+", ln.strip()):
                    break
                if "|" in ln and ln.strip().startswith("|"):
                    break
                if re.match(r"^(\*{3,}|-{3,}|_{3,})\s*$", ln.strip()):
                    break
                para_lines.append(ln)
                i += 1
            if para_lines:
                self.blocks.append({"type": "paragraph", "text": " ".join(para_lines)})

    def _parse_table(self, lines: list[str]):
        """Parse markdown table lines into a structured block."""
        rows = []
        for line in lines:
            # Skip separator row (---, :---:, etc.)
            if re.match(r"^\|[\s\-:]+\|$", line.strip()):
                continue
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            rows.append(cells)

        if not rows:
            return

        # Normalize column count
        max_cols = max(len(r) for r in rows)
        for r in rows:
            while len(r) < max_cols:
                r.append("")

        self.blocks.append({"type": "table", "rows": rows})


# ─── Inline text parser ──────────────────────────────────────────────────────

def parse_inline(text: str) -> list[dict]:
    """Parse inline markdown formatting (bold, italic, code, math) into runs."""
    runs: list[dict] = []
    # Pattern: **bold**, *italic*, `code`, $math$, ~~strikethrough~~
    pattern = re.compile(
        r"(?P<bold>\*\*(.+?)\*\*)"
        r"|(?P<italic>\*(.+?)\*)"
        r"|(?P<code>`(.+?)`)"
        r"|(?P<math>\$([^$]+?)\$)"
        r"|(?P<strike>~~(.+?)~~)"
    )

    last_end = 0
    for m in pattern.finditer(text):
        # Add plain text before this match
        if m.start() > last_end:
            plain = text[last_end:m.start()]
            if plain:
                runs.append({"text": plain, "bold": False, "italic": False, "code": False, "math": False})

        if m.group("bold"):
            runs.append({"text": m.group(2), "bold": True, "italic": False, "code": False, "math": False})
        elif m.group("italic"):
            runs.append({"text": m.group(4), "italic": True, "bold": False, "code": False, "math": False})
        elif m.group("code"):
            runs.append({"text": m.group(6), "code": True, "bold": False, "italic": False, "math": False})
        elif m.group("math"):
            runs.append({"text": m.group(8), "math": True, "bold": False, "italic": False, "code": False})
        elif m.group("strike"):
            runs.append({"text": m.group(10), "bold": False, "italic": False, "code": False, "math": False, "strike": True})

        last_end = m.end()

    # Remaining plain text
    if last_end < len(text):
        remaining = text[last_end:]
        if remaining:
            runs.append({"text": remaining, "bold": False, "italic": False, "code": False, "math": False})

    if not runs:
        runs.append({"text": text, "bold": False, "italic": False, "code": False, "math": False})

    return runs


# ─── DOCX builder ─────────────────────────────────────────────────────────────

class DocxBuilder:
    """Build a DOCX document from parsed markdown blocks."""

    def __init__(self, style_config: dict, math_converter: MathConverter = None, mermaid_renderer: MermaidRenderer = None):
        self.style_config = style_config
        self.doc = Document()
        self.math = math_converter or MathConverter()
        self.mermaid = mermaid_renderer or MermaidRenderer()
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles based on the selected preset."""
        cfg = self.style_config

        # Default paragraph style
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = cfg["body_font"]
        font.size = Pt(cfg["body_size"])
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pf = style.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = cfg["line_spacing"]

        # Set Vietnamese-compatible font for East Asian fallback
        style.element.rPr.rFonts.set(qn("w:eastAsia"), cfg["body_font"])

        # Heading styles
        for level in range(1, 7):
            hname = f"Heading {level}"
            if hname in self.doc.styles:
                hs = self.doc.styles[hname]
                hf = hs.font
                hf.name = cfg["heading_font"]
                hf.size = Pt(cfg["heading_sizes"].get(level, 11))
                hf.bold = True
                hf.color.rgb = cfg["heading_color"]
                hs.element.rPr.rFonts.set(qn("w:eastAsia"), cfg["heading_font"])
                hs.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
                hs.paragraph_format.space_after = Pt(6)

    def _add_inline_runs(self, paragraph, text: str, base_size: int = None):
        """Add inline-formatted runs to a paragraph."""
        cfg = self.style_config
        runs = parse_inline(text)
        for r in runs:
            # ── Inline math: try OMML first, fallback to text ──
            if r.get("math"):
                latex = r["text"]
                omml_elem = self.math.latex_to_omml(latex)
                if omml_elem is not None:
                    # Insert OMML element directly into paragraph
                    paragraph._element.append(omml_elem)
                    continue
                # Fallback: render as styled text
                run = paragraph.add_run(latex)
                run.font.name = "Cambria Math"
                run.italic = True
                run.font.color.rgb = RGBColor(0x00, 0x50, 0xA0)
                run.font.size = Pt(base_size or cfg["body_size"])
                continue

            run = paragraph.add_run(r["text"])
            run.font.name = cfg["body_font"]
            run.font.size = Pt(base_size or cfg["body_size"])

            if r.get("bold"):
                run.bold = True
            if r.get("italic"):
                run.italic = True
            if r.get("code"):
                run.font.name = cfg["code_font"]
                run.font.size = Pt((base_size or cfg["body_size"]) - 1)
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
                # Add shading for inline code
                shading = run.element.get_or_add_rPr()
                shd = shading.makeelement(qn("w:shd"), {
                    qn("w:val"): "clear",
                    qn("w:color"): "auto",
                    qn("w:fill"): "F0F0F0",
                })
                shading.append(shd)
            if r.get("strike"):
                run.font.strike = True

    def add_block(self, block: dict):
        """Add a parsed markdown block to the document."""
        btype = block["type"]

        if btype == "heading":
            self._add_heading(block)

        elif btype == "paragraph":
            self._add_paragraph(block)

        elif btype == "code":
            self._add_code(block)

        elif btype == "math_block":
            self._add_math_block(block)

        elif btype == "mermaid":
            self._add_mermaid(block)

        elif btype == "table":
            self._add_table(block)

        elif btype == "ul":
            self._add_list(block, ordered=False)

        elif btype == "ol":
            self._add_list(block, ordered=True)

        elif btype == "hr":
            self._add_hr()

    def _add_heading(self, block: dict):
        level = block["level"]
        text = block["text"]
        # Remove trailing hashes (### heading ###)
        text = re.sub(r"\s*#+\s*$", "", text).strip()
        hname = f"Heading {min(level, 6)}"
        para = self.doc.add_paragraph(style=hname)
        self._add_inline_runs(para, text, base_size=self.style_config["heading_sizes"].get(level, 11))

    def _add_paragraph(self, block: dict):
        para = self.doc.add_paragraph()
        self._add_inline_runs(para, block["text"])

    def _add_code(self, block: dict):
        cfg = self.style_config
        code = block["content"]
        lang = block.get("lang", "")

        # Add language label if present
        if lang:
            label = self.doc.add_paragraph()
            label_run = label.add_run(f"[{lang}]")
            label_run.font.name = cfg["code_font"]
            label_run.font.size = Pt(9)
            label_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            label.paragraph_format.space_after = Pt(0)

        # Add code content
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.left_indent = Cm(1)

        for line in code.split("\n"):
            run = para.add_run(line)
            run.font.name = cfg["code_font"]
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
            # Add line break within paragraph
            if line != code.split("\n")[-1]:
                para.add_run().add_break()

        # Add shading to code block
        pPr = para._element.get_or_add_pPr()
        shd = pPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "F5F5F5",
        })
        pPr.append(shd)

    def _add_math_block(self, block: dict):
        """Add a display math block ($$...$$) as OMML equation."""
        cfg = self.style_config
        content = block["content"]

        # Try OMML conversion for each line of the math block
        lines = [l.strip() for l in content.split("\n") if l.strip()]

        if self.math.available and lines:
            # Create an oMathPara (display math paragraph) for the block
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(8)

            for line in lines:
                omml_elem = self.math.latex_to_omml(line)
                if omml_elem is not None:
                    para._element.append(omml_elem)
                else:
                    # Fallback for failed conversion
                    run = para.add_run(line)
                    run.font.name = "Cambria Math"
                    run.font.size = Pt(cfg["body_size"])
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x00, 0x50, 0xA0)
        else:
            # Fallback: render as styled text
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(8)

            for idx, line in enumerate(lines):
                run = para.add_run(line)
                run.font.name = "Cambria Math"
                run.font.size = Pt(cfg["body_size"])
                run.italic = True
                run.font.color.rgb = RGBColor(0x00, 0x50, 0xA0)
                if idx < len(lines) - 1:
                    para.add_run().add_break()

    def _add_table(self, block: dict):
        cfg = self.style_config
        rows = block["rows"]
        if not rows:
            return

        num_cols = len(rows[0])
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = ""
                para = cell.paragraphs[0]

                if i == 0:
                    # Header row
                    run = para.add_run(cell_text)
                    run.bold = True
                    run.font.name = cfg["heading_font"]
                    run.font.size = Pt(cfg["body_size"] - 1)
                    run.font.color.rgb = cfg["table_header_fg"]
                    # Header background
                    shading = cell._element.get_or_add_tcPr()
                    shd = shading.makeelement(qn("w:shd"), {
                        qn("w:val"): "clear",
                        qn("w:color"): "auto",
                        qn("w:fill"): str(cfg["table_header_bg"]),
                    })
                    shading.append(shd)
                else:
                    self._add_inline_runs(para, cell_text, base_size=cfg["body_size"] - 1)

        # Add spacing after table
        self.doc.add_paragraph()

    def _add_list(self, block: dict, ordered: bool = False):
        cfg = self.style_config
        items = block["items"]
        for idx, item in enumerate(items):
            indent = item.get("indent", 0)
            text = item["text"]

            para = self.doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1.27 + indent * 1.27)
            para.paragraph_format.space_after = Pt(2)

            if ordered:
                prefix = f"{idx + 1}. "
            else:
                prefix = "• "

            run = para.add_run(prefix)
            run.font.name = cfg["body_font"]
            run.font.size = Pt(cfg["body_size"])
            run.bold = True

            self._add_inline_runs(para, text)

    def _add_mermaid(self, block: dict):
        """Render a Mermaid diagram and insert it as an image in the document."""
        cfg = self.style_config
        content = block["content"]

        # Try to render mermaid to PNG
        png_path = self.mermaid.render(content) if self.mermaid.available else None

        if png_path and Path(png_path).exists():
            # Add image to document
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)

            # Calculate width to fit page (keep aspect ratio)
            try:
                if HAS_PIL:
                    img = PILImage.open(png_path)
                    img_width, img_height = img.size
                    # Max width: 15cm (A4 usable width ~16cm, leave margin)
                    max_width_cm = 15.0
                    if img_width > 0:
                        aspect = img_height / img_width
                        doc_width = min(max_width_cm, img_width / 96 * 2.54)  # px to cm
                        doc_height = doc_width * aspect
                    else:
                        doc_width = max_width_cm
                        doc_height = max_width_cm * 0.5
                else:
                    doc_width = Inches(6)
                    doc_height = None

                run = para.add_run()
                run.add_picture(png_path, width=Cm(doc_width) if HAS_PIL else Inches(6))
            except Exception as e:
                print(f"  ⚠️ Failed to insert mermaid image: {e}", file=sys.stderr)
                # Fallback: add as code block
                self._add_code({"content": content, "lang": "mermaid"})
        else:
            # Fallback: render as styled code block with a label
            label = self.doc.add_paragraph()
            label_run = label.add_run("[Sơ đồ Mermaid]")
            label_run.font.name = cfg["heading_font"]
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
            label_run.bold = True
            label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label.paragraph_format.space_after = Pt(2)

            self._add_code({"content": content, "lang": "mermaid"})

    def _add_hr(self):
        """Add a horizontal rule as a thin border paragraph."""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)
        # Add a thin line using border
        pPr = para._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): "CCCCCC",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    def save(self, output_path: str):
        """Save the document to a file."""
        self.doc.save(output_path)


# ─── Main ─────────────────────────────────────────────────────────────────────

def convert_md_to_docx(input_path: str, output_path: str = None, font_size: int = None, style_name: str = "classic"):
    """Convert a Markdown file to DOCX.

    Args:
        input_path: Path to the input Markdown file.
        output_path: Path for the output DOCX file. Defaults to same name as input with .docx extension.
        font_size: Override body font size (pt).
        style_name: Style preset name (classic, modern, academic).
    """
    input_file = Path(input_path)
    if not input_file.exists():
        print(f"ERROR: Input file not found: {input_path}")
        sys.exit(1)

    if output_path is None:
        output_path = str(input_file.with_suffix(".docx"))

    # Read markdown
    md_text = input_file.read_text(encoding="utf-8")

    # Select style
    cfg = STYLES.get(style_name, STYLE_CLASSIC).copy()
    if font_size:
        cfg["body_size"] = font_size

    # Parse and build
    parser = MarkdownParser(md_text)
    math_converter = MathConverter()
    mermaid_renderer = MermaidRenderer()
    builder = DocxBuilder(cfg, math_converter=math_converter, mermaid_renderer=mermaid_renderer)

    for block in parser.blocks:
        builder.add_block(block)

    builder.save(output_path)
    math_status = "✅ OMML" if math_converter.available else "⚠️ text fallback"
    mermaid_status = "✅ mmdc" if mermaid_renderer.available else "⚠️ no mmdc"
    print(f"✅ Converted: {input_path} → {output_path}")
    print(f"   Blocks parsed: {len(parser.blocks)}")
    print(f"   Style: {style_name} | Font: {cfg['body_font']} {cfg['body_size']}pt")
    print(f"   Math engine: {math_status}")
    print(f"   Mermaid engine: {mermaid_status}")


def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown files to DOCX format",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python scripts/md_to_docx.py README.md
  python scripts/md_to_docx.py docs/report.md output.docx
  python scripts/md_to_docx.py docs/report.md --style modern
  python scripts/md_to_docx.py docs/report.md --font 14 --style academic
        """,
    )
    parser.add_argument("input", help="Input Markdown file path")
    parser.add_argument("output", nargs="?", default=None, help="Output DOCX file path (default: same name as input)")
    parser.add_argument("--font", type=int, default=None, help="Body font size in pt (default: style-specific)")
    parser.add_argument("--style", choices=list(STYLES.keys()), default="classic", help="Document style preset (default: classic)")

    args = parser.parse_args()
    convert_md_to_docx(args.input, args.output, args.font, args.style)


if __name__ == "__main__":
    main()