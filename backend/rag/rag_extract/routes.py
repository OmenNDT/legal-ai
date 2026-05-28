"""Flask Blueprint cho RAG Extract — tích hợp vào legal-ai.

Các endpoint tương thích với RAG_Extract gốc nhưng chạy trên Flask.
Auth: X-API-Key header (giống RAG_Extract gốc).
"""
import json
import logging
import os
from pathlib import Path
from typing import Optional

from flask import Blueprint, request, jsonify, Response, g

from backend.rag_extract.auth import require_api_key, require_admin
from backend.rag_extract.database import get_db, init_db
from backend.rag_extract.models.rag_document import RagDocument
from backend.rag_extract.services.rag_query_service import RagQueryService

logger = logging.getLogger(__name__)

rag_bp = Blueprint("rag_extract", __name__, url_prefix="/api/rag-extract")

# Lazy-init RAG service
_rag_service: Optional[RagQueryService] = None


def _get_rag_service() -> RagQueryService:
    global _rag_service
    if _rag_service is None:
        _rag_service = RagQueryService()
    return _rag_service


# ── Health ──
@rag_bp.get("/health")
def health():
    return jsonify({"status": "ok", "module": "rag_extract"})


# ── Stats ──
@rag_bp.get("/stats")
@require_api_key()
def get_stats():
    db = next(get_db())
    total = db.query(RagDocument).count()
    ready = db.query(RagDocument).filter(RagDocument.status == "ready").count()
    failed = db.query(RagDocument).filter(RagDocument.status == "failed").count()
    return jsonify({
        "total_documents": total,
        "ready_documents": ready,
        "failed_documents": failed,
    })


# ── Query ──
@rag_bp.post("/query")
@require_api_key()
def query_documents():
    body = request.get_json(force=True) or {}
    question = body.get("question", "")
    n_results = body.get("n_results", 5)

    if not question:
        return jsonify({"error": "Missing 'question' field"}), 400

    try:
        svc = _get_rag_service()
        # Note: async query cần event loop trong Flask
        import asyncio
        result = asyncio.run(svc.query(question, n_results=n_results))
        return jsonify(result)
    except Exception as e:
        logger.error(f"RAG query failed: {e}")
        return jsonify({"error": str(e)}), 500


# ── List Documents ──
@rag_bp.get("/documents")
@require_api_key()
def list_documents():
    db = next(get_db())
    skip = int(request.args.get("skip", 0))
    limit = int(request.args.get("limit", 100))
    docs = db.query(RagDocument).offset(skip).limit(limit).all()
    return jsonify({
        "items": [
            {
                "id": d.id,
                "filename": d.filename,
                "display_name": d.display_name,
                "standard_code": d.standard_code,
                "status": d.status,
                "processor": d.processor,
                "language": d.language,
                "created_at": d.created_at.isoformat() if d.created_at else None,
            }
            for d in docs
        ],
        "total": db.query(RagDocument).count(),
    })


# ── Get Document ──
@rag_bp.get("/documents/<int:doc_id>")
@require_api_key()
def get_document(doc_id: int):
    db = next(get_db())
    doc = db.query(RagDocument).filter(RagDocument.id == doc_id).first()
    if not doc:
        return jsonify({"error": "Document not found"}), 404
    return jsonify({
        "id": doc.id,
        "filename": doc.filename,
        "display_name": doc.display_name,
        "standard_code": doc.standard_code,
        "status": doc.status,
        "processor": doc.processor,
        "language": doc.language,
        "processed_content": doc.processed_content,
        "created_at": doc.created_at.isoformat() if doc.created_at else None,
    })


# ── Vector Store Stats ──
@rag_bp.get("/index/stats")
@require_api_key()
def index_stats():
    try:
        svc = _get_rag_service()
        return jsonify(svc.get_stats())
    except Exception as e:
        logger.error(f"Index stats failed: {e}")
        return jsonify({"error": str(e)}), 500


# ── Upload Document ──
@rag_bp.post("/upload")
@require_api_key()
def upload_document():
    """Upload one or more documents for RAG processing.

    Accepts multipart/form-data with:
      - files: one or more PDF/DOCX/TXT files
      - display_name (optional): friendly name
      - standard_code (optional): e.g., "ASME B31.3"
      - language (optional): 'en', 'vi', 'mixed' (default: 'vi')
    """
    if "files" not in request.files:
        return jsonify({"error": "No files provided. Use 'files' field."}), 400

    files = request.files.getlist("files")
    display_name = request.form.get("display_name", "")
    standard_code = request.form.get("standard_code", "")
    language = request.form.get("language", "vi")

    if not files or all(f.filename == "" for f in files):
        return jsonify({"error": "No files selected"}), 400

    ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".xlsx", ".xls"}
    uploaded = []
    errors = []

    db = next(get_db())

    for f in files:
        filename = f.filename
        if not filename:
            continue

        ext = Path(filename).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            errors.append({"filename": filename, "error": f"Unsupported file type: {ext}"})
            continue

        # Save file to upload directory
        from backend.rag_extract.config import UPLOAD_DIR
        safe_name = f"{Path(filename).stem}_{id(f)}{ext}"
        file_path = UPLOAD_DIR / safe_name
        f.save(str(file_path))

        # Get file size
        file_size = file_path.stat().st_size

        # Create DB record
        doc = RagDocument(
            filename=filename,
            display_name=display_name or Path(filename).stem,
            standard_code=standard_code or None,
            file_size_bytes=file_size,
            file_path=str(file_path.relative_to(UPLOAD_DIR.parent.parent)),
            language=language,
            status="processing",
            processor=None,
            uploaded_by=getattr(g, "rag_user", None) and g.rag_user.username or "api_user",
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)

        uploaded.append({
            "id": doc.id,
            "filename": doc.filename,
            "display_name": doc.display_name,
            "status": doc.status,
            "file_size_bytes": doc.file_size_bytes,
        })

        logger.info(f"Uploaded document: {filename} -> id={doc.id}")

    return jsonify({
        "uploaded": uploaded,
        "errors": errors,
        "total_uploaded": len(uploaded),
        "total_errors": len(errors),
    }), 201 if uploaded else 400


# ── Delete Document ──
@rag_bp.delete("/documents/<int:doc_id>")
@require_api_key(admin_only=True)
def delete_document(doc_id: int):
    """Delete a document and its associated files."""
    db = next(get_db())
    doc = db.query(RagDocument).filter(RagDocument.id == doc_id).first()
    if not doc:
        return jsonify({"error": "Document not found"}), 404

    # Delete file from disk
    from backend.rag_extract.config import DATA_DIR
    file_path = DATA_DIR / doc.file_path
    try:
        if file_path.exists():
            file_path.unlink()
            logger.info(f"Deleted file: {file_path}")
    except Exception as e:
        logger.warning(f"Could not delete file {file_path}: {e}")

    # Delete from vector store
    try:
        svc = _get_rag_service()
        svc._store.delete_by_doc_id(doc.id)
        logger.info(f"Deleted vector store entries for doc {doc.id}")
    except Exception as e:
        logger.warning(f"Could not delete vector store entries: {e}")

    # Delete DB record
    db.delete(doc)
    db.commit()

    return jsonify({"message": f"Document {doc_id} deleted", "id": doc_id})


# ── Process Document (trigger indexing) ──
@rag_bp.post("/documents/<int:doc_id>/process")
@require_api_key()
def process_document(doc_id: int):
    """Trigger processing/indexing for a document."""
    db = next(get_db())
    doc = db.query(RagDocument).filter(RagDocument.id == doc_id).first()
    if not doc:
        return jsonify({"error": "Document not found"}), 404

    if doc.status == "ready":
        return jsonify({"message": "Document already processed", "id": doc_id})

    try:
        from backend.rag_extract.config import DATA_DIR
        file_path = DATA_DIR / doc.file_path

        if not file_path.exists():
            doc.status = "failed"
            doc.error_message = f"File not found: {doc.file_path}"
            db.commit()
            return jsonify({"error": "File not found on disk"}), 404

        # Read file content
        content = _extract_text(file_path, doc.filename)
        if not content:
            doc.status = "failed"
            doc.error_message = "Could not extract text from file"
            db.commit()
            return jsonify({"error": "Could not extract text"}), 400

        doc.processed_content = content
        doc.status = "ready"
        doc.total_chunks = len(content) // 500  # rough estimate
        db.commit()

        # Index into vector store
        try:
            svc = _get_rag_service()
            num_chunks = svc.index_document(doc)
            doc.total_chunks = num_chunks
            db.commit()
        except Exception as e:
            logger.warning(f"Vector indexing failed for doc {doc_id}: {e}")
            # Still mark as ready since content is extracted

        return jsonify({
            "id": doc.id,
            "filename": doc.filename,
            "status": doc.status,
            "total_chunks": doc.total_chunks,
        })
    except Exception as e:
        logger.error(f"Process document failed: {e}")
        doc.status = "failed"
        doc.error_message = str(e)
        db.commit()
        return jsonify({"error": str(e)}), 500


def _extract_text(file_path: Path, filename: str) -> str:
    """Extract text from uploaded file based on extension."""
    ext = Path(filename).suffix.lower()

    if ext == ".txt" or ext == ".md":
        return file_path.read_text(encoding="utf-8", errors="replace")

    if ext == ".pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(file_path))
            text = "\n\n".join(page.get_text() for page in doc)
            doc.close()
            return text
        except ImportError:
            logger.warning("PyMuPDF not installed, trying pdfplumber")
            try:
                import pdfplumber
                with pdfplumber.open(str(file_path)) as pdf:
                    return "\n\n".join(p.extract_text() or "" for p in pdf.pages)
            except ImportError:
                logger.error("No PDF library available")
                return ""

    if ext in (".docx", ".doc"):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(file_path))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            logger.error("python-docx not installed")
            return ""

    # Fallback: try reading as text
    try:
        return file_path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return ""


# ── Init DB endpoint (admin only, one-time setup) ──
@rag_bp.post("/init-db")
@require_admin()
def init_database():
    try:
        init_db()
        return jsonify({"message": "Database initialized", "tables": ["rag_documents"]})
    except Exception as e:
        logger.error(f"Init DB failed: {e}")
        return jsonify({"error": str(e)}), 500
