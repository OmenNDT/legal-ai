import logging
from dataclasses import dataclass
from pathlib import Path
import io

logger = logging.getLogger(__name__)

@dataclass
class RawDocument:
    filename: str
    source_path: str
    raw_text: str
    file_format: str

class DocumentParser:
    _FORMAT_MAP = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".html": "html",
        ".htm": "html",
    }

    def detect_format(self, filename: str) -> str:
        return self._FORMAT_MAP.get(Path(filename).suffix.lower(), "text")

    def parse_pdf(self, content: bytes) -> str:
        from pdfminer.high_level import extract_text
        return extract_text(io.BytesIO(content)) or ""

    def parse_docx(self, content: bytes) -> str:
        from docx import Document
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def parse_html(self, content: bytes) -> str:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(content, "html.parser")
        for tag in soup(["script", "style", "header", "footer", "nav"]):
            tag.decompose()
        return soup.get_text(separator="\n")

    def parse(self, content: bytes, fmt: str) -> str:
        if fmt == "pdf":
            return self.parse_pdf(content)
        if fmt == "docx":
            return self.parse_docx(content)
        if fmt == "html":
            return self.parse_html(content)
        return content.decode("utf-8", errors="replace")

class DocumentIngester:
    def __init__(self, spark=None):
        self._spark = spark
        self._parser = DocumentParser()

    def ingest(self, source_path: str) -> list:
        if self._spark is not None:
            return self._ingest_spark(source_path)
        return self._ingest_local(source_path)

    def _ingest_local(self, source_path: str) -> list:
        path = Path(source_path)
        files = list(path.rglob("*")) if path.is_dir() else [path]
        documents = []
        for f in files:
            if not f.is_file():
                continue
            fmt = self._parser.detect_format(f.name)
            try:
                content = f.read_bytes()
                text = self._parser.parse(content, fmt)
                if text.strip():
                    documents.append(RawDocument(
                        filename=f.name,
                        source_path=str(f),
                        raw_text=text,
                        file_format=fmt,
                    ))
                    logger.info("[Ingest] OK: %s (%s, %d chars)", f.name, fmt, len(text))
                else:
                    logger.warning("[Ingest] Empty: %s", f.name)
            except Exception as e:
                logger.error("[Ingest] FAILED: %s — %s", f.name, e, exc_info=True)
        logger.info("[Ingest] Total: %d doc(s)", len(documents))
        return documents

    def _ingest_spark(self, source_path: str) -> list:
        binary_rdd = self._spark.sparkContext.binaryFiles(source_path)

        def _parse_partition(records):
            import io
            from pathlib import Path

            fmt_map = {
                ".pdf": "pdf", ".docx": "docx", ".doc": "docx",
                ".html": "html", ".htm": "html",
            }

            for path_str, content in records:
                filename = Path(path_str).name
                fmt = fmt_map.get(Path(filename).suffix.lower(), "text")
                try:
                    if fmt == "pdf":
                        from pdfminer.high_level import extract_text
                        text = extract_text(io.BytesIO(content)) or ""
                    elif fmt == "docx":
                        from docx import Document
                        doc = Document(io.BytesIO(content))
                        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    elif fmt == "html":
                        from bs4 import BeautifulSoup
                        soup = BeautifulSoup(content, "html.parser")
                        for tag in soup(["script", "style", "header", "footer", "nav"]):
                            tag.decompose()
                        text = soup.get_text(separator="\n")
                    else:
                        text = content.decode("utf-8", errors="replace")
                    if text.strip():
                        yield (filename, path_str, text, fmt)
                except Exception:
                    pass

        return [
            RawDocument(filename=r[0], source_path=r[1], raw_text=r[2], file_format=r[3])
            for r in binary_rdd.mapPartitions(_parse_partition).collect()
        ]
